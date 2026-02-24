"""
Exotel SOW Generator - Streamlit UI
AI-powered Statement of Work generator using Gemini 3 Flash
"""

import streamlit as st
import google.generativeai as genai
import json
import tempfile
import os
from datetime import datetime
from io import BytesIO
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
import numpy as np
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Page config
st.set_page_config(
    page_title="Exotel SOW Generator",
    page_icon="📄",
    layout="wide"
)

# Custom CSS
st.markdown("""
<style>
    .main-header {
        font-size: 2.5rem;
        font-weight: bold;
        color: #1E3A8A;
        margin-bottom: 0.5rem;
    }
    .sub-header {
        font-size: 1.1rem;
        color: #6B7280;
        margin-bottom: 2rem;
    }
    .success-box {
        padding: 1rem;
        border-radius: 0.5rem;
        background-color: #D1FAE5;
        border: 1px solid #10B981;
    }
    .info-box {
        padding: 1rem;
        border-radius: 0.5rem;
        background-color: #DBEAFE;
        border: 1px solid #3B82F6;
    }
</style>
""", unsafe_allow_html=True)

# Module catalog
MODULE_CATALOG = {
    "IVR": {
        "IVR-001": "Welcome Message",
        "IVR-002": "Language Selection",
        "IVR-003": "Main Menu",
        "IVR-007": "DTMF Input Capture",
        "IVR-010": "PIN Verification",
        "IVR-011": "OTP Generation & Validation",
        "IVR-015": "Consent Capture",
        "IVR-019": "Self-Service Balance Inquiry",
        "IVR-020": "Order Status Check",
        "IVR-023": "Callback Request",
        "IVR-029": "Time-Based Routing",
        "IVR-033": "Appointment Booking"
    },
    "Integration": {
        "INT-001": "REST API Outbound",
        "INT-002": "Webhook Configuration",
        "INT-004": "Salesforce Integration",
        "INT-007": "Zoho CRM Integration",
        "INT-008": "Freshdesk Integration",
        "INT-013": "Payment Gateway Integration",
        "INT-014": "SMS Gateway Integration",
        "INT-017": "Custom CRM Integration",
        "INT-021": "Screen Pop",
        "INT-023": "Post-Call Data Push"
    },
    "Blaster": {
        "BLA-002": "Outbound Campaign Setup",
        "BLA-003": "Auto Callback",
        "BLA-004": "Collections Campaign",
        "BLA-007": "Click-to-Dial",
        "BLA-010": "Progressive Dialer",
        "BLA-011": "Predictive Dialer",
        "BLA-018": "Retry Logic Configuration"
    },
    "Queue": {
        "QUE-001": "Basic Queue",
        "QUE-002": "Skill-Based Routing",
        "QUE-003": "Priority Queue",
        "QUE-007": "Sticky Agent",
        "QUE-012": "Callback from Queue",
        "QUE-019": "Multi-Queue Agent",
        "QUE-020": "Supervisor Monitor"
    },
    "Data": {
        "DAT-001": "CDR Export",
        "DAT-007": "Custom Reports",
        "DAT-009": "Real-Time Dashboard",
        "DAT-010": "Survey Responses",
        "DAT-011": "Agent Performance Reports"
    }
}

def init_gemini(api_key):
    """Initialize Gemini client"""
    genai.configure(api_key=api_key)
    return genai.GenerativeModel('gemini-3-flash-preview')

def extract_requirements(model, transcript):
    """Extract structured requirements from transcript"""
    prompt = f"""You are an expert at analyzing call transcripts for Exotel, a cloud communications platform.

## Task
Extract structured requirements from this call transcript.

## Call Transcript:
{transcript}

## Output Format:
Return a valid JSON object with:
{{
    "customer": "Customer Name",
    "project": "Project Name",
    "industry": "Industry type",
    "attendees": [{{"name": "Name", "role": "Role"}}],
    "requirements": ["Requirement 1", "Requirement 2"],
    "modules_needed": ["Module descriptions"],
    "integrations": ["Integration needs"],
    "timeline": "Expected timeline",
    "notes": ["Additional notes"]
}}

Return ONLY the JSON, no other text."""

    response = model.generate_content(prompt)
    text = response.text

    if "```json" in text:
        text = text.split("```json")[1].split("```")[0]
    elif "```" in text:
        text = text.split("```")[1].split("```")[0]

    return json.loads(text.strip())

def generate_sow_content(model, requirements):
    """Generate full SOW content in Exotel format"""
    prompt = f"""You are an expert technical writer for Exotel/Ameyo cloud communications platform.

## Task
Generate a detailed Statement of Work (SOW) based on these requirements, following the exact Exotel SOW format.

## Requirements:
{json.dumps(requirements, indent=2)}

## Available Modules:
{json.dumps(MODULE_CATALOG, indent=2)}

## Output Format:
Return a JSON object with this EXACT structure:
{{
    "customer": "Customer Name (short code like TENB)",
    "customer_full_name": "Full Customer Company Name",
    "project": "Project Name",
    "version": "1.0.0",
    "date": "{datetime.now().strftime('%d-%m-%Y')}",
    "ticket_id": "DCA{datetime.now().strftime('%Y%m%d')}XXXXXX",
    "created_by": "PS Team",
    "business_goals": [
        {{
            "sno": "1",
            "use_case": "Detailed description of the business goal or use case the customer wants to achieve",
            "deliverables": "What Exotel team will deliver to meet this goal"
        }}
    ],
    "prerequisites": [
        "Prerequisite 1 that customer team needs to provide",
        "Prerequisite 2"
    ],
    "deliverables": [
        {{
            "title": "Deliverable Title (e.g., Call Transfer from VoiceBot to ECC)",
            "description": "Detailed multi-paragraph description of what will be implemented, how it works technically, the flow, and expected behavior. Be specific about SIP, API, routing logic, etc."
        }}
    ],
    "deployment": [
        "Deployment will be done in a single setup.",
        "Specific deployment details about campaigns, instances, etc."
    ],
    "notes": [
        "Important note about scope validity",
        "Note about changes outside SOW being treated as separate service ticket",
        "This Scope of Work is valid for the next 2 months from the last revision date."
    ],
    "assumptions": [
        "Exotel/Ameyo Setup will be commissioned at the center as per the setup project plan.",
        "The client team will provide Access while testing and deployment.",
        "The client team's support might be required for customization activities.",
        "The environmental prerequisites such as availability of structured data center, power, hardware, and network stability are assumed to be provided by the client."
    ],
    "account_manager": {{
        "name": "Account Manager Name",
        "email": "email@exotel.com"
    }},
    "escalation_scoping": [
        {{"level": "1", "name": "Name", "phone": "+91 XXXXXXXXXX", "email": "email@exotel.com"}},
        {{"level": "2", "name": "Name", "phone": "+91 XXXXXXXXXX", "email": "email@exotel.com"}}
    ],
    "escalation_deployment": [
        {{"level": "1", "name": "Name", "phone": "+91 XXXXXXXXXX", "email": "email@exotel.com"}},
        {{"level": "2", "name": "Name", "phone": "+91 XXXXXXXXXX", "email": "email@exotel.com"}},
        {{"level": "3", "name": "Name", "phone": "+91 XXXXXXXXXX", "email": "email@exotel.com"}}
    ]
}}

IMPORTANT:
- Be very detailed in the deliverables description - write multiple paragraphs explaining the technical implementation
- Use professional language similar to: "Client team wants to enable...", "The system will...", "During the process..."
- Include technical details about SIP, API, routing, campaigns where relevant
- Make assumptions realistic for a cloud communications deployment

Return ONLY the JSON."""

    response = model.generate_content(prompt)
    text = response.text

    if "```json" in text:
        text = text.split("```json")[1].split("```")[0]
    elif "```" in text:
        text = text.split("```")[1].split("```")[0]

    return json.loads(text.strip())

def generate_flowchart_structure(model, requirements):
    """Generate flowchart structure"""
    prompt = f"""You are an expert at creating IVR and call flow diagrams.

## Task
Generate a flowchart structure for these requirements.

## Requirements:
{json.dumps(requirements, indent=2)}

## Output Format:
Return a JSON object with:
{{
    "title": "Flowchart Title",
    "nodes": [
        {{
            "id": "node_1",
            "type": "start",
            "label": "Start"
        }},
        {{
            "id": "node_2",
            "type": "process",
            "label": "Process Step"
        }}
    ],
    "edges": [
        {{
            "from": "node_1",
            "to": "node_2",
            "label": ""
        }}
    ]
}}

Node types: start, end, process, decision, api, queue, disconnect

Return ONLY the JSON."""

    response = model.generate_content(prompt)
    text = response.text

    if "```json" in text:
        text = text.split("```json")[1].split("```")[0]
    elif "```" in text:
        text = text.split("```")[1].split("```")[0]

    return json.loads(text.strip())

def generate_drawio_xml(flowchart):
    """Convert flowchart structure to draw.io XML"""
    node_styles = {
        "start": "ellipse;whiteSpace=wrap;html=1;fillColor=#d5e8d4;strokeColor=#82b366;",
        "end": "ellipse;whiteSpace=wrap;html=1;fillColor=#f8cecc;strokeColor=#b85450;",
        "process": "rounded=1;whiteSpace=wrap;html=1;fillColor=#dae8fc;strokeColor=#6c8ebf;",
        "decision": "rhombus;whiteSpace=wrap;html=1;fillColor=#fff2cc;strokeColor=#d6b656;",
        "api": "rounded=1;whiteSpace=wrap;html=1;fillColor=#e1d5e7;strokeColor=#9673a6;",
        "queue": "rounded=1;whiteSpace=wrap;html=1;fillColor=#d5e8fc;strokeColor=#6c8ebf;",
        "disconnect": "ellipse;whiteSpace=wrap;html=1;fillColor=#f8cecc;strokeColor=#b85450;"
    }

    nodes = flowchart.get("nodes", [])
    edges = flowchart.get("edges", [])

    cells = []
    node_positions = {}

    y_pos = 40
    for i, node in enumerate(nodes):
        node_id = node.get("id", f"node_{i}")
        node_type = node.get("type", "process")
        label = node.get("label", "")
        style = node_styles.get(node_type, node_styles["process"])

        width = 120 if node_type != "decision" else 100
        height = 60 if node_type != "decision" else 80
        x_pos = 300

        node_positions[node_id] = {"x": x_pos, "y": y_pos, "w": width, "h": height}

        cells.append(f'''      <mxCell id="{node_id}" value="{label}" style="{style}" vertex="1" parent="1">
        <mxGeometry x="{x_pos}" y="{y_pos}" width="{width}" height="{height}" as="geometry"/>
      </mxCell>''')

        y_pos += height + 60

    for i, edge in enumerate(edges):
        from_id = edge.get("from", "")
        to_id = edge.get("to", "")
        label = edge.get("label", "")

        cells.append(f'''      <mxCell id="edge_{i}" value="{label}" style="edgeStyle=orthogonalEdgeStyle;rounded=0;orthogonalLoop=1;jettySize=auto;html=1;" edge="1" parent="1" source="{from_id}" target="{to_id}">
        <mxGeometry relative="1" as="geometry"/>
      </mxCell>''')

    xml = f'''<?xml version="1.0" encoding="UTF-8"?>
<mxfile host="app.diagrams.net" modified="{datetime.now().isoformat()}" agent="Exotel SOW Generator" version="1.0">
  <diagram name="Flow" id="flow-diagram">
    <mxGraphModel dx="1200" dy="800" grid="1" gridSize="10" guides="1" tooltips="1" connect="1" arrows="1" fold="1" page="1" pageScale="1" pageWidth="827" pageHeight="1169">
      <root>
        <mxCell id="0"/>
        <mxCell id="1" parent="0"/>
{chr(10).join(cells)}
      </root>
    </mxGraphModel>
  </diagram>
</mxfile>'''

    return xml

def generate_flowchart_image(flowchart):
    """Generate a visual flowchart image using matplotlib"""
    nodes = flowchart.get("nodes", [])
    edges = flowchart.get("edges", [])
    title = flowchart.get("title", "IVR Flow")

    colors = {
        "start": "#d5e8d4",
        "end": "#f8cecc",
        "process": "#dae8fc",
        "decision": "#fff2cc",
        "api": "#e1d5e7",
        "queue": "#d5e8fc",
        "disconnect": "#f8cecc"
    }

    edge_colors = {
        "start": "#82b366",
        "end": "#b85450",
        "process": "#6c8ebf",
        "decision": "#d6b656",
        "api": "#9673a6",
        "queue": "#6c8ebf",
        "disconnect": "#b85450"
    }

    fig_height = max(8, len(nodes) * 1.5)
    fig, ax = plt.subplots(1, 1, figsize=(10, fig_height))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, fig_height)
    ax.set_aspect('equal')
    ax.axis('off')

    ax.text(5, fig_height - 0.5, title, fontsize=16, fontweight='bold',
            ha='center', va='top')

    node_positions = {}
    y_pos = fig_height - 1.5
    x_center = 5

    for i, node in enumerate(nodes):
        node_id = node.get("id", f"node_{i}")
        node_type = node.get("type", "process")
        label = node.get("label", "")
        color = colors.get(node_type, colors["process"])
        edge_color = edge_colors.get(node_type, edge_colors["process"])

        node_positions[node_id] = (x_center, y_pos)

        if node_type in ["start", "end", "disconnect"]:
            ellipse = mpatches.Ellipse((x_center, y_pos), 3, 0.8,
                                        facecolor=color, edgecolor=edge_color, linewidth=2)
            ax.add_patch(ellipse)
        elif node_type == "decision":
            diamond = plt.Polygon([(x_center, y_pos + 0.5), (x_center + 1.5, y_pos),
                                   (x_center, y_pos - 0.5), (x_center - 1.5, y_pos)],
                                  facecolor=color, edgecolor=edge_color, linewidth=2)
            ax.add_patch(diamond)
        else:
            rect = FancyBboxPatch((x_center - 1.5, y_pos - 0.4), 3, 0.8,
                                   boxstyle="round,pad=0.05,rounding_size=0.2",
                                   facecolor=color, edgecolor=edge_color, linewidth=2)
            ax.add_patch(rect)

        ax.text(x_center, y_pos, label, fontsize=9, ha='center', va='center',
                wrap=True, fontweight='medium')

        y_pos -= 1.3

    for edge in edges:
        from_id = edge.get("from", "")
        to_id = edge.get("to", "")
        edge_label = edge.get("label", "")

        if from_id in node_positions and to_id in node_positions:
            from_pos = node_positions[from_id]
            to_pos = node_positions[to_id]

            ax.annotate("", xy=(to_pos[0], to_pos[1] + 0.45),
                       xytext=(from_pos[0], from_pos[1] - 0.45),
                       arrowprops=dict(arrowstyle="->", color="#666666", lw=1.5))

            if edge_label:
                mid_y = (from_pos[1] + to_pos[1]) / 2
                ax.text(x_center + 0.3, mid_y, edge_label, fontsize=8,
                       ha='left', va='center', color="#666666")

    legend_y = 0.8
    legend_items = [
        ("Start/End", "#d5e8d4", "#82b366"),
        ("Process", "#dae8fc", "#6c8ebf"),
        ("Decision", "#fff2cc", "#d6b656"),
        ("API Call", "#e1d5e7", "#9673a6"),
        ("Queue", "#d5e8fc", "#6c8ebf"),
    ]

    for i, (name, fc, ec) in enumerate(legend_items):
        rect = FancyBboxPatch((0.3, legend_y - i * 0.4), 0.4, 0.25,
                               boxstyle="round,pad=0.02",
                               facecolor=fc, edgecolor=ec, linewidth=1)
        ax.add_patch(rect)
        ax.text(0.9, legend_y - i * 0.4 + 0.12, name, fontsize=7, va='center')

    plt.tight_layout()

    buf = BytesIO()
    plt.savefig(buf, format='pdf', bbox_inches='tight', dpi=150)
    buf.seek(0)
    pdf_data = buf.getvalue()

    buf_png = BytesIO()
    plt.savefig(buf_png, format='png', bbox_inches='tight', dpi=150)
    buf_png.seek(0)
    png_data = buf_png.getvalue()

    plt.close()

    return pdf_data, png_data

def set_cell_shading(cell, color):
    """Set cell background color"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def generate_sow_docx(sow):
    """Generate a Word document from SOW content in Exotel format"""
    doc = Document()

    # Set page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(0.94)

    # ===== TITLE PAGE =====
    # Main Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Scope of Work")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0, 51, 102)

    # Customer Name
    customer_title = doc.add_paragraph()
    customer_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = customer_title.add_run(sow.get('customer', 'Customer'))
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0, 51, 102)

    doc.add_paragraph()

    # Document Properties Table
    props_table = doc.add_table(rows=6, cols=2)
    props_table.style = 'Table Grid'

    props_data = [
        ('Document Properties', ''),
        ('Creation Date', sow.get('date', datetime.now().strftime('%d-%m-%Y'))),
        ('Current Draft', sow.get('version', 'v1.0.0')),
        ('Modification Date', '-'),
        ('Ticket ID', sow.get('ticket_id', f"DCA{datetime.now().strftime('%Y%m%d')}000000")),
        ('Created By', sow.get('created_by', 'PS Team'))
    ]

    for i, (label, value) in enumerate(props_data):
        row = props_table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = value
        if i == 0:
            row.cells[0].merge(row.cells[1])
            row.cells[0].paragraphs[0].runs[0].bold = True
            set_cell_shading(row.cells[0], 'D9E2F3')
        else:
            row.cells[0].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()

    # Copyright
    copyright_table = doc.add_table(rows=1, cols=1)
    copyright_table.style = 'Table Grid'
    cell = copyright_table.rows[0].cells[0]
    cell.text = f"© {datetime.now().year} eXotel, All rights reserved.   CONFIDENTIAL"
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_shading(cell, 'D9E2F3')

    # Page break
    doc.add_page_break()

    # ===== DOCUMENT HISTORY =====
    doc.add_paragraph("Document History").runs[0].bold = True

    history_table = doc.add_table(rows=2, cols=4)
    history_table.style = 'Table Grid'

    headers = ['Version', 'Creation date', 'Created By', 'Summary of changes']
    for i, header in enumerate(headers):
        cell = history_table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, 'D9E2F3')

    history_table.rows[1].cells[0].text = sow.get('version', '1.0.0')
    history_table.rows[1].cells[1].text = sow.get('date', datetime.now().strftime('%d-%m-%Y'))
    history_table.rows[1].cells[2].text = sow.get('created_by', 'PS Team')
    history_table.rows[1].cells[3].text = 'Define Scope of Work'

    doc.add_paragraph()

    # Table of Contents placeholder
    doc.add_paragraph("Table of Contents").runs[0].bold = True
    doc.add_paragraph()

    # ===== 1. BUSINESS GOAL VS DELIVERABLES =====
    h1 = doc.add_heading('1. Business Goal Vs Deliverables', level=1)

    business_goals = sow.get('business_goals', [])
    if business_goals:
        bg_table = doc.add_table(rows=1, cols=3)
        bg_table.style = 'Table Grid'

        headers = ['S.No', 'Business Goal / Use Case', 'Deliverables']
        for i, header in enumerate(headers):
            cell = bg_table.rows[0].cells[i]
            cell.text = header
            cell.paragraphs[0].runs[0].bold = True
            set_cell_shading(cell, 'D9E2F3')

        for goal in business_goals:
            row = bg_table.add_row()
            row.cells[0].text = str(goal.get('sno', '1'))
            row.cells[1].text = goal.get('use_case', '')
            row.cells[2].text = goal.get('deliverables', '')

    doc.add_paragraph()

    # ===== 2. PREREQUISITES & LICENSES =====
    doc.add_heading('2. Prerequisites & Licenses', level=1)

    customer_name = sow.get('customer', 'Customer')
    doc.add_paragraph(f'Below are the dependencies on the "{customer_name}" team:')
    doc.add_paragraph(f'{customer_name} team will be required to provide the below mentioned details:')

    for prereq in sow.get('prerequisites', []):
        doc.add_paragraph(prereq, style='List Bullet')

    doc.add_paragraph()

    # ===== 3. DETAILS OF DELIVERABLES =====
    doc.add_heading('3. Details of Deliverables', level=1)

    deliverables = sow.get('deliverables', [])
    for idx, deliverable in enumerate(deliverables, 1):
        doc.add_heading(f'3.{idx}. {deliverable.get("title", "Deliverable")}', level=2)

        # Add description paragraphs
        desc = deliverable.get('description', '')
        paragraphs = desc.split('\n\n') if '\n\n' in desc else [desc]
        for para in paragraphs:
            if para.strip():
                doc.add_paragraph(para.strip())

    doc.add_paragraph()

    # ===== 4. DEPLOYMENT & UAT INSTANCE =====
    doc.add_heading('4. Deployment & UAT Instance', level=1)

    for deployment in sow.get('deployment', ['Deployment will be done in a single setup.']):
        doc.add_paragraph(deployment)

    doc.add_paragraph()

    # ===== 5. NOTES =====
    doc.add_heading('5. Notes', level=1)

    for note in sow.get('notes', []):
        doc.add_paragraph(note)

    doc.add_paragraph()

    # ===== 6. ASSUMPTIONS =====
    doc.add_heading('6. Assumptions', level=1)

    for assumption in sow.get('assumptions', []):
        doc.add_paragraph(assumption)

    doc.add_paragraph()

    # ===== 7. ANNEXURE =====
    doc.add_heading('7. Annexure', level=1)

    # 7.1 Account Manager
    doc.add_heading('7.1. Account Manager from Exotel', level=2)
    am = sow.get('account_manager', {})
    doc.add_paragraph(f"             Name: {am.get('name', 'TBD')}")
    doc.add_paragraph(f"             Email: {am.get('email', 'TBD')}")

    # 7.2 Escalation Matrix for Scoping
    doc.add_heading('7.2. Escalation Matrix for Scoping', level=2)

    esc_scoping = sow.get('escalation_scoping', [])
    if esc_scoping:
        esc_table = doc.add_table(rows=1, cols=4)
        esc_table.style = 'Table Grid'

        headers = ['Level', 'Name', 'Phone', 'Email ID']
        for i, header in enumerate(headers):
            cell = esc_table.rows[0].cells[i]
            cell.text = header
            cell.paragraphs[0].runs[0].bold = True
            set_cell_shading(cell, 'D9E2F3')

        for esc in esc_scoping:
            row = esc_table.add_row()
            row.cells[0].text = str(esc.get('level', ''))
            row.cells[1].text = esc.get('name', '')
            row.cells[2].text = esc.get('phone', '')
            row.cells[3].text = esc.get('email', '')

    doc.add_paragraph()

    # 7.3 Escalation Matrix for Deployment
    doc.add_heading('7.3. Escalation Matrix for Deployment', level=2)

    esc_deploy = sow.get('escalation_deployment', [])
    if esc_deploy:
        esc_table2 = doc.add_table(rows=1, cols=4)
        esc_table2.style = 'Table Grid'

        for i, header in enumerate(headers):
            cell = esc_table2.rows[0].cells[i]
            cell.text = header
            cell.paragraphs[0].runs[0].bold = True
            set_cell_shading(cell, 'D9E2F3')

        for esc in esc_deploy:
            row = esc_table2.add_row()
            row.cells[0].text = str(esc.get('level', ''))
            row.cells[1].text = esc.get('name', '')
            row.cells[2].text = esc.get('phone', '')
            row.cells[3].text = esc.get('email', '')

    doc.add_paragraph()

    # 7.4 SOW Approval Procedure
    doc.add_heading('7.4. SOW Approval Procedure', level=2)
    doc.add_paragraph("Dear Customer, in order to approve the Scope, kindly click on the approve button on your Salesforce portal.")

    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run('Generated by Exotel SOW Generator')
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(128, 128, 128)

    # Save to BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

def format_sow_markdown(sow):
    """Format SOW as markdown for display"""
    md = f"""# Scope of Work - {sow.get('customer', 'Customer')}

**Version:** {sow.get('version', '1.0.0')}
**Date:** {sow.get('date', datetime.now().strftime('%d-%m-%Y'))}
**Ticket ID:** {sow.get('ticket_id', 'N/A')}
**Created By:** {sow.get('created_by', 'PS Team')}

---

## 1. Business Goal Vs Deliverables

"""
    for goal in sow.get('business_goals', []):
        md += f"### {goal.get('sno', '1')}. Use Case\n"
        md += f"**Business Goal:** {goal.get('use_case', 'N/A')}\n\n"
        md += f"**Deliverables:** {goal.get('deliverables', 'N/A')}\n\n"

    md += "\n## 2. Prerequisites & Licenses\n\n"
    for prereq in sow.get('prerequisites', []):
        md += f"- {prereq}\n"

    md += "\n## 3. Details of Deliverables\n\n"
    for idx, deliv in enumerate(sow.get('deliverables', []), 1):
        md += f"### 3.{idx}. {deliv.get('title', 'Deliverable')}\n\n"
        md += f"{deliv.get('description', 'N/A')}\n\n"

    md += "\n## 4. Deployment & UAT Instance\n\n"
    for dep in sow.get('deployment', []):
        md += f"- {dep}\n"

    md += "\n## 5. Notes\n\n"
    for note in sow.get('notes', []):
        md += f"- {note}\n"

    md += "\n## 6. Assumptions\n\n"
    for assumption in sow.get('assumptions', []):
        md += f"- {assumption}\n"

    md += "\n## 7. Annexure\n\n"
    am = sow.get('account_manager', {})
    md += f"### 7.1. Account Manager\n- Name: {am.get('name', 'TBD')}\n- Email: {am.get('email', 'TBD')}\n\n"

    md += "### 7.2. Escalation Matrix for Scoping\n"
    md += "| Level | Name | Phone | Email |\n|-------|------|-------|-------|\n"
    for esc in sow.get('escalation_scoping', []):
        md += f"| {esc.get('level', '')} | {esc.get('name', '')} | {esc.get('phone', '')} | {esc.get('email', '')} |\n"

    md += "\n### 7.3. Escalation Matrix for Deployment\n"
    md += "| Level | Name | Phone | Email |\n|-------|------|-------|-------|\n"
    for esc in sow.get('escalation_deployment', []):
        md += f"| {esc.get('level', '')} | {esc.get('name', '')} | {esc.get('phone', '')} | {esc.get('email', '')} |\n"

    return md

# Main UI
st.markdown('<p class="main-header">📄 Exotel SOW Generator</p>', unsafe_allow_html=True)
st.markdown('<p class="sub-header">AI-powered Statement of Work generator using Gemini 3 Flash</p>', unsafe_allow_html=True)

api_key = None

try:
    api_key = st.secrets.get("GEMINI_API_KEY")
except:
    pass

if not api_key:
    api_key = os.environ.get("GEMINI_API_KEY")

with st.sidebar:
    st.header("⚙️ Configuration")

    if not api_key:
        api_key = st.text_input(
            "Gemini API Key",
            type="password",
            help="Get your API key from Google AI Studio"
        )
    else:
        st.success("✅ API Key configured")

    st.divider()

    st.header("📊 Module Catalog")
    with st.expander("View Available Modules"):
        for category, modules in MODULE_CATALOG.items():
            st.subheader(category)
            for mod_id, mod_name in modules.items():
                st.text(f"{mod_id}: {mod_name}")

    st.divider()

    st.markdown("""
    ### About
    This tool generates professional SOW documents from call transcripts using AI.

    **Powered by:**
    - Gemini 3 Flash
    - 1,800+ historical SOWs
    - Exotel SOW Template
    """)

if not api_key:
    st.warning("⚠️ API Key not configured. Add GEMINI_API_KEY to Streamlit secrets or enter it in the sidebar.")
    st.stop()

try:
    model = init_gemini(api_key)
except Exception as e:
    st.error(f"Failed to initialize Gemini: {e}")
    st.stop()

st.header("📤 Upload Call Transcript")

uploaded_file = st.file_uploader(
    "Upload a call transcript or meeting notes",
    type=["txt", "pdf", "md"],
    help="Supported formats: TXT, PDF, MD"
)

transcript_text = st.text_area(
    "Or paste your transcript here:",
    height=200,
    placeholder="Paste your call transcript, meeting notes, or requirements here..."
)

if st.button("🚀 Generate SOW", type="primary", use_container_width=True):
    transcript = ""

    if uploaded_file:
        if uploaded_file.type == "text/plain":
            transcript = uploaded_file.read().decode("utf-8")
        elif uploaded_file.type == "application/pdf":
            try:
                import PyPDF2
                pdf_reader = PyPDF2.PdfReader(uploaded_file)
                transcript = "\n".join([page.extract_text() for page in pdf_reader.pages])
            except ImportError:
                st.error("PDF support requires PyPDF2. Please paste text instead.")
                st.stop()
        else:
            transcript = uploaded_file.read().decode("utf-8")
    elif transcript_text:
        transcript = transcript_text
    else:
        st.warning("Please upload a file or paste your transcript")
        st.stop()

    if len(transcript) < 50:
        st.warning("Transcript seems too short. Please provide more details.")
        st.stop()

    with st.spinner("🔍 Analyzing transcript..."):
        try:
            requirements = extract_requirements(model, transcript)
            st.success("✅ Requirements extracted!")
        except Exception as e:
            st.error(f"Failed to extract requirements: {e}")
            st.stop()

    with st.spinner("📝 Generating SOW..."):
        try:
            sow_content = generate_sow_content(model, requirements)
            st.success("✅ SOW generated!")
        except Exception as e:
            st.error(f"Failed to generate SOW: {e}")
            st.stop()

    with st.spinner("📊 Creating flowchart..."):
        try:
            flowchart = generate_flowchart_structure(model, requirements)
            drawio_xml = generate_drawio_xml(flowchart)
            st.success("✅ Flowchart created!")
        except Exception as e:
            st.warning(f"Flowchart generation failed: {e}")
            flowchart = None
            drawio_xml = None

    st.divider()
    st.header("📋 Generated SOW")

    tab1, tab2, tab3 = st.tabs(["📄 SOW Document", "🔍 Requirements", "📊 Flowchart"])

    with tab1:
        st.markdown(format_sow_markdown(sow_content))

        col1, col2, col3 = st.columns(3)
        with col1:
            docx_data = generate_sow_docx(sow_content)
            st.download_button(
                "📥 Download SOW (Word)",
                data=docx_data,
                file_name=f"SOW_{sow_content.get('customer', 'Customer').replace(' ', '_')}_v{sow_content.get('version', '1.0.0')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        with col2:
            st.download_button(
                "📥 Download SOW (JSON)",
                data=json.dumps(sow_content, indent=2),
                file_name=f"SOW_{sow_content.get('customer', 'Customer').replace(' ', '_')}_v{sow_content.get('version', '1.0.0')}.json",
                mime="application/json"
            )
        with col3:
            st.download_button(
                "📥 Download SOW (Markdown)",
                data=format_sow_markdown(sow_content),
                file_name=f"SOW_{sow_content.get('customer', 'Customer').replace(' ', '_')}_v{sow_content.get('version', '1.0.0')}.md",
                mime="text/markdown"
            )

    with tab2:
        st.subheader("Extracted Requirements")
        st.json(requirements)

    with tab3:
        if flowchart and drawio_xml:
            st.subheader("Flow Diagram")

            try:
                pdf_data, png_data = generate_flowchart_image(flowchart)

                st.image(png_data, caption="IVR Flow Diagram", use_container_width=True)

                col1, col2, col3 = st.columns(3)
                with col1:
                    st.download_button(
                        "📥 Download PDF",
                        data=pdf_data,
                        file_name=f"{sow_content.get('project', 'Flow').replace(' ', '_')}_flow.pdf",
                        mime="application/pdf"
                    )
                with col2:
                    st.download_button(
                        "📥 Download PNG",
                        data=png_data,
                        file_name=f"{sow_content.get('project', 'Flow').replace(' ', '_')}_flow.png",
                        mime="image/png"
                    )
                with col3:
                    st.download_button(
                        "📥 Download .drawio",
                        data=drawio_xml,
                        file_name=f"{sow_content.get('project', 'Flow').replace(' ', '_')}_flow.drawio",
                        mime="application/xml"
                    )
            except Exception as e:
                st.warning(f"Could not generate visual flowchart: {e}")
                st.json(flowchart)
                st.download_button(
                    "📥 Download Flowchart (.drawio)",
                    data=drawio_xml,
                    file_name=f"{sow_content.get('project', 'Flow').replace(' ', '_')}_flow.drawio",
                    mime="application/xml"
                )

            with st.expander("View Flow Structure (JSON)"):
                st.json(flowchart)
        else:
            st.info("Flowchart not available")

st.divider()
st.markdown("""
<div style="text-align: center; color: #6B7280; font-size: 0.9rem;">
    Built with ❤️ by Exotel PS Team | Powered by Gemini 3 Flash
</div>
""", unsafe_allow_html=True)
