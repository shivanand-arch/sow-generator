"""
SOW Generator v2 — Output Generators (DOCX + Google Doc)
"""

import os
import subprocess
import json
import re
from pathlib import Path
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from config import OUTPUT_DIR, GOOGLE_AUTH_SCRIPT


# ============================================================
# DOCX Generator
# ============================================================

class DocxGenerator:
    """Generate formatted DOCX from Markdown SOW content."""

    def generate(self, markdown_text, spec, output_path=None):
        """Generate DOCX from assembled Markdown."""
        if output_path is None:
            OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
            slug = (spec.customer.name or "Unknown").replace(" ", "_")[:25]
            ts = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_path = OUTPUT_DIR / f"SOW_{slug}_{ts}.docx"

        doc = Document()

        # Set default font
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(11)

        # Process markdown line by line
        lines = markdown_text.split("\n")
        i = 0
        while i < len(lines):
            line = lines[i]

            # Headers
            if line.startswith("# ") and not line.startswith("## "):
                p = doc.add_heading(line[2:].strip(), level=0)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif line.startswith("## "):
                doc.add_heading(line[3:].strip(), level=1)
            elif line.startswith("### "):
                doc.add_heading(line[4:].strip(), level=2)
            elif line.startswith("#### "):
                doc.add_heading(line[5:].strip(), level=3)

            # Tables
            elif "|" in line and i + 1 < len(lines) and "---" in lines[i + 1]:
                table_lines = [line]
                j = i + 1
                while j < len(lines) and "|" in lines[j]:
                    table_lines.append(lines[j])
                    j += 1
                self._add_table(doc, table_lines)
                i = j
                continue

            # Bullet lists
            elif line.strip().startswith("- ") or line.strip().startswith("* "):
                text = line.strip()[2:].strip()
                p = doc.add_paragraph(style="List Bullet")
                self._add_formatted_text(p, text)
            elif re.match(r'^\d+\.\s', line.strip()):
                text = re.sub(r'^\d+\.\s', '', line.strip())
                p = doc.add_paragraph(style="List Number")
                self._add_formatted_text(p, text)

            # Horizontal rule
            elif line.strip() == "---":
                doc.add_paragraph("_" * 60)

            # Block quotes
            elif line.strip().startswith("> "):
                text = line.strip()[2:]
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(1)
                run = p.add_run(text)
                run.italic = True
                run.font.color.rgb = RGBColor(100, 100, 100)

            # Regular paragraphs
            elif line.strip():
                p = doc.add_paragraph()
                self._add_formatted_text(p, line.strip())

            i += 1

        doc.save(str(output_path))
        return str(output_path)

    def _add_table(self, doc, table_lines):
        """Parse Markdown table and add to document."""
        # Parse header
        headers = [c.strip() for c in table_lines[0].split("|") if c.strip()]
        if not headers:
            return

        # Parse rows (skip separator line)
        rows = []
        for line in table_lines[2:]:  # Skip header and separator
            if "---" in line:
                continue
            cells = [c.strip() for c in line.split("|") if c.strip()]
            if cells:
                rows.append(cells)

        if not rows:
            return

        # Create table
        table = doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.style = "Light Grid Accent 1"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header row
        for j, header in enumerate(headers):
            cell = table.rows[0].cells[j]
            cell.text = header
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(10)

        # Data rows
        for i, row in enumerate(rows):
            for j, cell_text in enumerate(row):
                if j < len(headers):
                    cell = table.rows[i + 1].cells[j]
                    cell.text = cell_text
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(10)

    def _add_formatted_text(self, paragraph, text):
        """Add text with basic Markdown formatting (bold, italic, code)."""
        # Simple regex-based formatting
        parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`|\[VERIFY\])', text)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            elif part.startswith("*") and part.endswith("*") and not part.startswith("**"):
                run = paragraph.add_run(part[1:-1])
                run.italic = True
            elif part.startswith("`") and part.endswith("`"):
                run = paragraph.add_run(part[1:-1])
                run.font.name = "Consolas"
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(180, 0, 0)
            elif part == "[VERIFY]":
                run = paragraph.add_run("[VERIFY]")
                run.bold = True
                run.font.color.rgb = RGBColor(255, 140, 0)
            elif part:
                paragraph.add_run(part)


# ============================================================
# Google Doc Generator
# ============================================================

class GoogleDocGenerator:
    """Create a Google Doc from Markdown SOW content via Docs API."""

    def _get_access_token(self):
        """Get OAuth access token using the auth helper script."""
        try:
            result = subprocess.run(
                ["bash", "-c", f"source {GOOGLE_AUTH_SCRIPT} && echo $GOOGLE_ACCESS_TOKEN"],
                capture_output=True, text=True, timeout=15,
            )
            token = result.stdout.strip()
            if token and len(token) > 20:
                return token
        except Exception as e:
            print(f"[GoogleDoc] Auth failed: {e}")
        return None

    def create(self, markdown_text, spec):
        """
        Create a Google Doc with the SOW content.
        Returns: Google Doc URL or None on failure.
        """
        token = self._get_access_token()
        if not token:
            print("[GoogleDoc] No access token — skipping Google Doc creation")
            return None

        customer = spec.customer.name or "Unknown"
        title = f"SOW — {customer} ECC Implementation ({datetime.now().strftime('%Y-%m-%d')})"

        try:
            import requests

            # Step 1: Create empty doc
            create_resp = requests.post(
                "https://docs.googleapis.com/v1/documents",
                headers={
                    "Authorization": f"Bearer {token}",
                    "Content-Type": "application/json",
                },
                json={"title": title},
                timeout=15,
            )
            create_resp.raise_for_status()
            doc_id = create_resp.json()["documentId"]

            # Step 2: Insert content as plain text (Docs API batch update)
            # Convert markdown to plain text for insertion (formatting via API is complex)
            plain_text = self._markdown_to_plaintext(markdown_text)

            requests_body = {
                "requests": [
                    {
                        "insertText": {
                            "location": {"index": 1},
                            "text": plain_text,
                        }
                    }
                ]
            }

            update_resp = requests.post(
                f"https://docs.googleapis.com/v1/documents/{doc_id}:batchUpdate",
                headers={
                    "Authorization": f"Bearer {token}",
                    "Content-Type": "application/json",
                },
                json=requests_body,
                timeout=30,
            )
            update_resp.raise_for_status()

            doc_url = f"https://docs.google.com/document/d/{doc_id}/edit"
            print(f"[GoogleDoc] Created: {doc_url}")
            return doc_url

        except Exception as e:
            print(f"[GoogleDoc] Creation failed: {e}")
            return None

    def _markdown_to_plaintext(self, markdown):
        """Basic Markdown → plain text conversion for Google Docs insertion."""
        text = markdown
        # Remove Markdown formatting but preserve structure
        text = re.sub(r'^#{1,4}\s+', '', text, flags=re.MULTILINE)  # Headers → plain
        text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)  # Bold
        text = re.sub(r'\*(.*?)\*', r'\1', text)  # Italic
        text = re.sub(r'`(.*?)`', r'\1', text)  # Code
        text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)  # Links
        return text
